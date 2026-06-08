"""Document loaders — turn raw bytes into a list of LlamaIndex Documents.

Dispatch by MIME type:
  - application/pdf  → pypdf
  - .docx            → python-docx
  - text/plain       → utf-8 decode

We deliberately use the low-level libraries directly (instead of
`llama_index.readers.file`) because we already have the bytes in memory —
no need for the LlamaIndex readers' temp-file roundtrip.
"""

from __future__ import annotations

import io
import logging
from typing import Any

logger = logging.getLogger(__name__)


class LoaderError(Exception):
    """Raised when the upload can't be parsed into usable text.

    Caller maps this to HTTP 422 with ``error_code = "PARSE_ERROR"``.
    """


def load_documents(file_bytes: bytes, mime_type: str, original_name: str) -> list[Any]:
    """Returns a list of :class:`llama_index.core.Document` objects.

    Empty files / no-text PDFs raise :class:`LoaderError`.
    """
    if not file_bytes:
        raise LoaderError("Uploaded file is empty.")

    if mime_type == "text/plain":
        return _load_text(file_bytes, original_name)

    if mime_type == "application/pdf":
        return _load_pdf(file_bytes, original_name)

    if mime_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return _load_docx(file_bytes, original_name)
    
    if mime_type == "application/msword":
        return _load_legacy_doc(file_bytes, original_name)
    

    raise LoaderError(f"Unsupported MIME type: {mime_type}")

def _load_legacy_doc(file_bytes: bytes, original_name: str) -> list[Any]:
    """Read a legacy binary ``.doc`` (pre-2007 OLE format).

    Pipeline:
        bytes (.doc)  ── LibreOffice headless ──►  .docx  ──►  _load_docx

    Why not call ``mammoth`` directly?  Mammoth only understands the
    OOXML / docx zip format.  The binary OLE ``.doc`` format produces
    ``Could not find file 'word/document.xml'`` — silent dead end.

    The on-disk conversion uses two temp dirs; both are wiped in the
    ``finally`` block on every code path (success, parse error, crash).
    """
    import shutil
    import tempfile
    from pathlib import Path

    from ..preprocessing.docx_processor import (
        DocxProcessingError,
        convert_doc_to_docx,
    )

    # Suffix MUST be .doc — LibreOffice picks the import filter by ext.
    src_dir = Path(tempfile.mkdtemp(prefix="ld3_doc_in_"))
    out_dir: Path | None = None
    try:
        # Use a sanitised stem so non-ASCII filenames don't break soffice
        # (LibreOffice is fussy about filename encoding on Windows).
        safe_stem = "input"
        src_path  = src_dir / f"{safe_stem}.doc"
        src_path.write_bytes(file_bytes)

        try:
            converted = convert_doc_to_docx(src_path)
        except DocxProcessingError as exc:
            raise LoaderError(f"DOC → DOCX conversion failed: {exc}") from exc

        out_dir     = converted.parent
        docx_bytes  = converted.read_bytes()
    finally:
        shutil.rmtree(src_dir, ignore_errors=True)
        if out_dir is not None:
            shutil.rmtree(out_dir, ignore_errors=True)

    if not docx_bytes:
        raise LoaderError("LibreOffice produced an empty .docx — corrupt source?")

    # Hand off to the existing docx loader (tables preserved via mammoth).
    return _load_docx(docx_bytes, original_name)
# ---------------------------------------------------------------------
#  Per-format loaders
# ---------------------------------------------------------------------

def _load_text(file_bytes: bytes, original_name: str) -> list[Any]:
    from llama_index.core import Document

    text = file_bytes.decode("utf-8", errors="replace")
    if not text.strip():
        raise LoaderError("TXT file has no readable content.")
    return [Document(text=text, metadata={"source": original_name})]


def _load_pdf(file_bytes: bytes, original_name: str) -> list[Any]:
    from llama_index.core import Document
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:
        raise LoaderError(f"PDF parse failed: {exc}") from exc

    pages: list[Any] = []
    for idx, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            logger.warning("pdf_page_extract_failed page=%s err=%s", idx, exc)
            text = ""
        if text.strip():
            pages.append(
                Document(
                    text=text,
                    metadata={"source": original_name, "page": idx + 1},
                )
            )

    if not pages:
        raise LoaderError("PDF contains no extractable text (likely scanned — OCR not enabled).")
    return pages


def _load_docx(file_bytes: bytes, original_name: str) -> list[Any]:
    from docx import Document as DocxDocument  # python-docx
    from llama_index.core import Document

    try:
        docx_obj = DocxDocument(io.BytesIO(file_bytes))
    except Exception as exc:
        raise LoaderError(f"DOCX parse failed: {exc}") from exc

    paragraphs = [p.text for p in docx_obj.paragraphs if p.text and p.text.strip()]
    if not paragraphs:
        raise LoaderError("DOCX has no readable text.")

    return [
        Document(
            text="\n".join(paragraphs),
            metadata={"source": original_name},
        )
    ]
